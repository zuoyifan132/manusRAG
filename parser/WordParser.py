from docx import Document
from parser.baseParser import BaseParser


class WordParser(BaseParser):
    def __init__(self, word_path):
        super().__init__(word_path)


class DocxParser(WordParser):
    def __init__(self, word_path):
        super().__init__(word_path)
    
    def read_content(self):
        """读取Word文档内容并存储到self.content中"""
        try:
            self.content = Document(self.path)
        except Exception as e:
            raise Exception(f"无法读取Word文档: {str(e)}")
    
    def extract_text(self):
        """从Word文档中提取文本"""
        if self.content is None:
            raise ValueError("内容未加载。请先调用read_content()")
        
        full_text = ""
        
        # 提取段落文本
        for para in self.content.paragraphs:
            if para.text:
                full_text += para.text + "\n"
        
        # 提取表格文本
        for table in self.content.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text:
                        row_text.append(cell.text)
                if row_text:
                    full_text += " | ".join(row_text) + "\n"
        
        return full_text


# 示例用法
if __name__ == "__main__":
    """示例：如何使用WordParser解析Word文档"""
    # 文档路径
    word_file_path = "/Users/evan/Desktop/rest/wzy/cv/Resume_RachelWang_fin.docx"
    
    try:
        # 创建DocxParser实例
        parser = DocxParser(word_file_path)
        
        # 读取文档内容
        parser.read_content()
        
        # 提取文本
        text = parser.extract_text()
        
        # 打印提取的文本
        print("从Word文档中提取的文本：")
        print("-" * 50)
        print(text)
        print("-" * 50)
        
        # 可以在这里继续处理提取的文本
        # 例如：存储到数据库，进行文本分析等
        
    except Exception as e:
        print(f"处理Word文档时出错: {str(e)}")

